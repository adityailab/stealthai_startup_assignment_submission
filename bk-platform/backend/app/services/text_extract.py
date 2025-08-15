#########try1#######first one
# import io
# from typing import Optional
# from pypdf import PdfReader
# from docx import Document as DocxDocument

# def io_bytes(b: bytes) -> io.BytesIO:
#     """Wrap bytes in a BytesIO object."""
#     bio = io.BytesIO()
#     bio.write(b)
#     bio.seek(0)
#     return bio

# def extract_text_from_bytes(data: bytes, filename: str, mime_type: str) -> Optional[str]:
#     """Extract text from PDF, DOCX, TXT, or MD files."""
#     name = filename.lower()
#     try:
#         # PDF
#         if name.endswith(".pdf") or mime_type == "application/pdf":
#             reader = PdfReader(io_bytes(data))
#             return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
#         # DOCX
#         if name.endswith(".docx") or mime_type in (
#             "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         ):
#             doc = DocxDocument(io_bytes(data))
#             return "\n".join(p.text for p in doc.paragraphs).strip()
#         # TXT / MD / Plain text
#         if name.endswith(".md") or name.endswith(".txt") or mime_type.startswith("text/"):
#             try:
#                 return data.decode("utf-8")
#             except UnicodeDecodeError:
#                 return data.decode("latin-1")
#     except Exception:
#         return None
#     return None


#######try2#####later one

# app/services/text_extract.py

from __future__ import annotations
import io
from typing import Optional
from pathlib import Path

# PDF text
from pypdf import PdfReader

# DOCX
try:
    import docx  # python-docx
except Exception:
    docx = None

# OCR stack
try:
    from pdf2image import convert_from_bytes  # requires poppler installed
    from PIL import Image
    import pytesseract
except Exception:
    convert_from_bytes = None
    Image = None
    pytesseract = None


TEXT_MIMES = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/json",
}


def _extract_pdf_text(data: bytes) -> str:
    """Try extracting selectable text from a PDF."""
    try:
        reader = PdfReader(io.BytesIO(data))
        chunks = []
        for page in reader.pages:
            txt = page.extract_text() or ""
            if txt:
                chunks.append(txt)
        return "\n".join(chunks).strip()
    except Exception:
        return ""


def _ocr_pdf(data: bytes) -> str:
    """OCR a scanned/image PDF using pdf2image + Tesseract."""
    if not (convert_from_bytes and pytesseract):
        return ""
    try:
        # 200–300 dpi is a good balance for OCR
        images = convert_from_bytes(data, dpi=250)  # list[PIL.Image]
        out = []
        for im in images:
            # Convert to grayscale improves OCR sometimes
            if hasattr(im, "convert"):
                im = im.convert("L")
            text = pytesseract.image_to_string(im, lang="eng") or ""
            if text:
                out.append(text)
        return "\n".join(out).strip()
    except Exception:
        return ""


def _ocr_image(data: bytes) -> str:
    """OCR a single image file (png/jpg/webp)."""
    if not (Image and pytesseract):
        return ""
    try:
        im = Image.open(io.BytesIO(data))
        if hasattr(im, "convert"):
            im = im.convert("L")
        return (pytesseract.image_to_string(im, lang="eng") or "").strip()
    except Exception:
        return ""


def _extract_docx(data: bytes) -> str:
    if not docx:
        return ""
    try:
        doc = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text).strip()
    except Exception:
        return ""


def extract_text_from_bytes(data: bytes, filename: str, mime: Optional[str]) -> str:
    """
    Best-effort text extraction:
    1) Text PDFs via pypdf
    2) If empty, OCR the PDF
    3) DOCX via python-docx
    4) Plain/Markdown/JSON as-is
    5) Image files via OCR
    """
    name = (filename or "").lower()
    m = (mime or "").lower()

    # PDF
    if m == "application/pdf" or name.endswith(".pdf"):
        text = _extract_pdf_text(data)
        if text and text.strip():
            return text
        # fall back to OCR for scanned PDFs
        return _ocr_pdf(data)

    # DOCX
    if m in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"} or name.endswith(".docx"):
        return _extract_docx(data)

    # Plain-ish
    if m in TEXT_MIMES or any(name.endswith(ext) for ext in (".txt", ".md", ".csv", ".json")):
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return ""

    # Images -> OCR
    if any(name.endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".webp")) or m.startswith("image/"):
        return _ocr_image(data)

    # Unknown
    return ""
